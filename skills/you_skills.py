# SPDX-License-Identifier: Elastic-2.0
# Copyright (c) 2026 Cabir Pekdemir. All rights reserved.
# Licensed under the Elastic License 2.0 — see LICENSE for details.

"""
OZY2 — YOU Tier Skills
Package: you

Includes: weather, web_search, news, currency, summarize,
          word_docx, excel_xlsx, notes, reminders
"""
from core.tools import register


def register_all():

    # ── Weather ───────────────────────────────────────────────────────────────
    @register(
        name="get_weather",
        description=(
            "Get current weather or forecast for any city. "
            "Use for: 'weather', 'is it raining', 'forecast this week'."
        ),
        params={
            "city":     {"type": "string",  "description": "City name (default: Istanbul)"},
            "forecast": {"type": "boolean", "description": "True for multi-day forecast"},
            "days":     {"type": "integer", "description": "Days ahead 1-7 (default 3)"},
        },
        package="you",
    )
    async def _weather(city: str = "Istanbul", forecast: bool = False, days: int = 3):
        import urllib.request, json
        # Geocode
        geo_url = f"https://geocoding-api.open-meteo.com/v1/search?name={urllib.parse.quote(city)}&count=1&language=en"
        import urllib.parse
        with urllib.request.urlopen(geo_url, timeout=8) as r:
            geo = json.loads(r.read())
        results = geo.get("results")
        if not results:
            return {"error": f"City '{city}' not found"}
        lat, lon = results[0]["latitude"], results[0]["longitude"]
        name     = results[0].get("name", city)

        WMO = {
            0: "Clear sky ☀️", 1: "Mainly clear 🌤️", 2: "Partly cloudy ⛅", 3: "Overcast ☁️",
            45: "Foggy 🌫️", 51: "Light drizzle 🌦️", 61: "Light rain 🌧️", 63: "Rain 🌧️",
            65: "Heavy rain 🌧️", 71: "Light snow 🌨️", 73: "Snow ❄️", 80: "Showers 🌦️",
            95: "Thunderstorm ⛈️", 99: "Heavy thunderstorm ⛈️",
        }

        if not forecast:
            url = (f"https://api.open-meteo.com/v1/forecast?latitude={lat}&longitude={lon}"
                   "&current=temperature_2m,weathercode,windspeed_10m,relative_humidity_2m"
                   "&timezone=auto")
            with urllib.request.urlopen(url, timeout=8) as r:
                data = json.loads(r.read())
            cur = data["current"]
            code = cur["weathercode"]
            return {
                "city": name, "temp": cur["temperature_2m"], "unit": "°C",
                "condition": WMO.get(code, f"Code {code}"),
                "wind_kmh": cur["windspeed_10m"],
                "humidity": cur["relative_humidity_2m"],
            }
        else:
            days = max(1, min(days, 7))
            url = (f"https://api.open-meteo.com/v1/forecast?latitude={lat}&longitude={lon}"
                   f"&daily=temperature_2m_max,temperature_2m_min,weathercode"
                   f"&timezone=auto&forecast_days={days}")
            with urllib.request.urlopen(url, timeout=8) as r:
                data = json.loads(r.read())
            d = data["daily"]
            return {
                "city": name,
                "forecast": [
                    {"date": d["time"][i],
                     "max": d["temperature_2m_max"][i],
                     "min": d["temperature_2m_min"][i],
                     "condition": WMO.get(d["weathercode"][i], "")}
                    for i in range(len(d["time"]))
                ]
            }

    # ── Web Search ────────────────────────────────────────────────────────────
    @register(
        name="web_search",
        description=(
            "Search the web. Use for current events, research, facts, "
            "anything that needs up-to-date information."
        ),
        params={
            "query":       {"type": "string",  "description": "Search query", "required": True},
            "max_results": {"type": "integer", "description": "Number of results (default 8)"},
            "news_only":   {"type": "boolean", "description": "Search news only"},
        },
        package="you",
    )
    async def _web_search(query: str, max_results: int = 8, news_only: bool = False):
        try:
            try:
                from ddgs import DDGS
            except ImportError:
                from duckduckgo_search import DDGS
            results = []
            with DDGS() as ddgs:
                fn = ddgs.news if news_only else ddgs.text
                for r in fn(query, max_results=max_results):
                    results.append({
                        "title":   r.get("title", ""),
                        "url":     r.get("href") or r.get("url", ""),
                        "snippet": r.get("body") or r.get("excerpt", ""),
                    })
            return {"ok": True, "query": query, "results": results}
        except Exception as e:
            return {"error": str(e), "query": query}

    # ── News ──────────────────────────────────────────────────────────────────
    @register(
        name="get_news",
        description=(
            "Get latest news headlines. Use for: 'news today', 'tech news', "
            "'sports headlines', 'what happened in X'."
        ),
        params={
            "topic":  {"type": "string",  "description": "Topic or keyword (optional, e.g. 'technology')"},
            "limit":  {"type": "integer", "description": "Number of articles (default 10)"},
        },
        package="you",
    )
    async def _news(topic: str = "", limit: int = 10):
        try:
            try:
                from ddgs import DDGS
            except ImportError:
                from duckduckgo_search import DDGS
            q = f"{topic} news" if topic else "top news today"
            results = []
            with DDGS() as ddgs:
                for r in ddgs.news(q, max_results=limit):
                    results.append({
                        "title":  r.get("title", ""),
                        "source": r.get("source", ""),
                        "date":   r.get("date", ""),
                        "url":    r.get("url", ""),
                        "body":   r.get("body", ""),
                    })
            return {"ok": True, "topic": topic or "general", "articles": results}
        except Exception as e:
            return {"error": str(e)}

    # ── Currency ──────────────────────────────────────────────────────────────
    @register(
        name="convert_currency",
        description=(
            "Convert between currencies or get live exchange rates. "
            "Use for: 'USD to EUR', '100 dollars in yen', 'exchange rate'."
        ),
        params={
            "from_currency": {"type": "string", "description": "Source currency code (e.g. USD)", "required": True},
            "to_currency":   {"type": "string", "description": "Target currency code (e.g. EUR)", "required": True},
            "amount":        {"type": "number", "description": "Amount to convert (default 1)"},
        },
        package="you",
    )
    async def _currency(from_currency: str, to_currency: str, amount: float = 1.0):
        import urllib.request, json
        base = from_currency.upper()
        target = to_currency.upper()
        url = f"https://open.er-api.com/v6/latest/{base}"
        try:
            with urllib.request.urlopen(url, timeout=8) as r:
                data = json.loads(r.read())
            if data.get("result") != "success":
                return {"error": f"Invalid currency: {base}"}
            rate = data["rates"].get(target)
            if rate is None:
                return {"error": f"Unknown currency: {target}"}
            return {
                "from": base, "to": target,
                "rate": rate,
                "amount": amount,
                "converted": round(amount * rate, 4),
                "updated": data.get("time_last_update_utc", ""),
            }
        except Exception as e:
            return {"error": str(e)}

    # ── Summarize ─────────────────────────────────────────────────────────────
    @register(
        name="summarize_text",
        description=(
            "Summarize long text or a webpage URL. "
            "Use for: 'summarize this', 'TL;DR', 'what does this article say', 'short version'."
        ),
        params={
            "text":   {"type": "string", "description": "Text content to summarize"},
            "url":    {"type": "string", "description": "URL to fetch and summarize (optional)"},
            "length": {"type": "string", "description": "short | medium | long (default medium)"},
        },
        package="you",
    )
    async def _summarize(text: str = "", url: str = "", length: str = "medium"):
        import urllib.request
        if url and not text:
            try:
                req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
                with urllib.request.urlopen(req, timeout=10) as r:
                    raw = r.read().decode("utf-8", errors="ignore")
                # Strip HTML tags roughly
                import re
                text = re.sub(r"<[^>]+>", " ", raw)
                text = re.sub(r"\s+", " ", text).strip()[:6000]
            except Exception as e:
                return {"error": f"Could not fetch URL: {e}"}
        if not text:
            return {"error": "Provide text or url parameter"}

        word_limits = {"short": 80, "medium": 150, "long": 280}
        limit = word_limits.get(length, 150)
        return {
            "source": url or "text",
            "length": length,
            "word_limit": limit,
            "text_preview": text[:200],
            "instruction": f"Please summarize the following text in ~{limit} words:\n\n{text[:4000]}",
        }

    # ── Word / DOCX ───────────────────────────────────────────────────────────
    @register(
        name="create_document",
        description=(
            "Create a Word (.docx) document. "
            "Use for: 'write a report', 'create a letter', 'make a Word file', 'draft a contract'."
        ),
        params={
            "title":    {"type": "string", "description": "Document title", "required": True},
            "content":  {"type": "string", "description": "Document body (markdown or plain text)"},
            "filename": {"type": "string", "description": "Output filename without extension"},
        },
        package="you",
    )
    async def _create_doc(title: str, content: str = "", filename: str = ""):
        try:
            from docx import Document
            from docx.shared import Pt
            import re, os
            from pathlib import Path

            doc = Document()
            doc.add_heading(title, 0)
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("- ") or line.startswith("* "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    doc.add_paragraph(line)

            fname = (filename or re.sub(r"[^\w\-]", "_", title)).rstrip("_") + ".docx"
            out_dir = Path.home() / "Downloads"
            out_dir.mkdir(exist_ok=True)
            path = str(out_dir / fname)
            doc.save(path)
            return {"ok": True, "file": path, "title": title}
        except ImportError:
            return {"error": "python-docx not installed. Run: pip install python-docx"}
        except Exception as e:
            return {"error": str(e)}

    # ── Excel / XLSX ──────────────────────────────────────────────────────────
    @register(
        name="create_spreadsheet",
        description=(
            "Create an Excel (.xlsx) spreadsheet. "
            "Use for: 'make a spreadsheet', 'export to Excel', 'create a table in xlsx'."
        ),
        params={
            "title":    {"type": "string", "description": "Sheet title / filename"},
            "headers":  {"type": "string", "description": "Comma-separated column headers"},
            "rows":     {"type": "string", "description": "JSON array of row arrays"},
            "filename": {"type": "string", "description": "Output filename without extension"},
        },
        package="you",
    )
    async def _create_xlsx(title: str = "Sheet1", headers: str = "",
                           rows: str = "[]", filename: str = ""):
        try:
            import openpyxl, json
            from pathlib import Path
            from openpyxl.styles import Font, PatternFill, Alignment
            import re

            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = title[:31]

            header_list = [h.strip() for h in headers.split(",") if h.strip()] if headers else []
            if header_list:
                ws.append(header_list)
                for cell in ws[1]:
                    cell.font = Font(bold=True, color="FFFFFF")
                    cell.fill = PatternFill("solid", fgColor="4f8ef7")
                    cell.alignment = Alignment(horizontal="center")

            row_data = json.loads(rows) if rows != "[]" else []
            for row in row_data:
                ws.append(row if isinstance(row, list) else list(row.values()))

            fname = (filename or re.sub(r"[^\w\-]", "_", title)).rstrip("_") + ".xlsx"
            out_dir = Path.home() / "Downloads"
            out_dir.mkdir(exist_ok=True)
            path = str(out_dir / fname)
            wb.save(path)
            return {"ok": True, "file": path, "rows": len(row_data), "columns": len(header_list)}
        except ImportError:
            return {"error": "openpyxl not installed. Run: pip install openpyxl"}
        except Exception as e:
            return {"error": str(e)}

    # ── Notes (Apple Notes / local) ───────────────────────────────────────────
    @register(
        name="add_note",
        description=(
            "Add a note. On macOS uses Apple Notes; otherwise saves locally. "
            "Use for: 'take a note', 'remember this idea', 'note that...'."
        ),
        params={
            "title":   {"type": "string", "description": "Note title"},
            "content": {"type": "string", "description": "Note content", "required": True},
            "folder":  {"type": "string", "description": "Folder / notebook name (default: OZY2)"},
        },
        package="you",
    )
    async def _add_note(content: str, title: str = "OZY Note", folder: str = "OZY2"):
        import sys, subprocess, json
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
        if sys.platform == "darwin":
            script = f'''
tell application "Notes"
    tell account "iCloud"
        if not (exists folder "{folder}") then
            make new folder with properties {{name:"{folder}"}}
        end if
        make new note at folder "{folder}" with properties {{name:"{title}", body:"{content}"}}
    end tell
end tell
'''
            try:
                subprocess.run(["osascript", "-e", script], timeout=10, check=True,
                               capture_output=True)
                return {"ok": True, "platform": "Apple Notes", "folder": folder, "title": title}
            except Exception:
                pass
        # Fallback: local file
        from pathlib import Path
        notes_dir = Path.home() / ".ozy2" / "notes"
        notes_dir.mkdir(parents=True, exist_ok=True)
        fname = notes_dir / f"{timestamp.replace(':', '-')}_{title[:30]}.txt"
        fname.write_text(f"# {title}\n{timestamp}\n\n{content}")
        return {"ok": True, "platform": "local", "file": str(fname)}

    @register(
        name="list_notes",
        description="List recent notes.",
        params={"limit": {"type": "integer", "description": "Max notes to return (default 10)"}},
        package="you",
    )
    async def _list_notes(limit: int = 10):
        import sys, subprocess
        from pathlib import Path
        if sys.platform == "darwin":
            script = f'''
tell application "Notes"
    set noteList to ""
    repeat with n in (first {limit} notes)
        set noteList to noteList & name of n & "\\n"
    end repeat
    return noteList
end tell
'''
            try:
                r = subprocess.run(["osascript", "-e", script], capture_output=True,
                                   text=True, timeout=10)
                notes = [l for l in r.stdout.strip().split("\n") if l]
                return {"ok": True, "platform": "Apple Notes", "notes": notes}
            except Exception:
                pass
        notes_dir = Path.home() / ".ozy2" / "notes"
        if notes_dir.exists():
            files = sorted(notes_dir.glob("*.txt"), reverse=True)[:limit]
            return {"ok": True, "platform": "local",
                    "notes": [f.stem for f in files]}
        return {"notes": []}

    # ── Reminders ─────────────────────────────────────────────────────────────
    @register(
        name="add_reminder",
        description=(
            "Set a reminder. On macOS uses Apple Reminders. "
            "Use for: 'remind me to...', 'set a reminder', 'don't let me forget'."
        ),
        params={
            "title":    {"type": "string", "description": "Reminder title", "required": True},
            "due_date": {"type": "string", "description": "Due date/time ISO or natural language"},
            "notes":    {"type": "string", "description": "Optional notes"},
            "list":     {"type": "string", "description": "Reminders list name (default: OZY2)"},
        },
        package="you",
    )
    async def _add_reminder(title: str, due_date: str = "", notes: str = "", list: str = "OZY2"):
        import sys, subprocess
        if sys.platform == "darwin":
            due_part = f', due date:date "{due_date}"' if due_date else ""
            note_part = f', body:"{notes}"' if notes else ""
            script = f'''
tell application "Reminders"
    if not (exists list "{list}") then
        make new list with properties {{name:"{list}"}}
    end if
    make new reminder at list "{list}" with properties {{name:"{title}"{due_part}{note_part}}}
end tell
'''
            try:
                subprocess.run(["osascript", "-e", script], timeout=10, check=True,
                               capture_output=True)
                return {"ok": True, "platform": "Apple Reminders", "title": title, "due": due_date}
            except Exception as e:
                return {"error": str(e)}
        return {"ok": False, "error": "Apple Reminders only supported on macOS"}

    @register(
        name="list_reminders",
        description="List upcoming reminders.",
        params={"list": {"type": "string", "description": "List name (default: all)"}},
        package="you",
    )
    async def _list_reminders(list: str = ""):
        import sys, subprocess
        if sys.platform == "darwin":
            list_filter = f'list "{list}"' if list else "every list"
            script = f'''
tell application "Reminders"
    set reminderList to ""
    repeat with r in (reminders of {list_filter} whose completed is false)
        set reminderList to reminderList & name of r & "\\n"
    end repeat
    return reminderList
end tell
'''
            try:
                r = subprocess.run(["osascript", "-e", script], capture_output=True,
                                   text=True, timeout=10)
                items = [l for l in r.stdout.strip().split("\n") if l]
                return {"ok": True, "reminders": items}
            except Exception as e:
                return {"error": str(e)}
        return {"error": "Apple Reminders only supported on macOS"}

    # ── Outfit of the Day ─────────────────────────────────────────────────────
    @register(
        name="outfit_of_day",
        description=(
            "Suggest an outfit for today based on weather and occasion. "
            "Use for: 'what should I wear', 'outfit for today', 'dress suggestion', 'bugün ne giysem'."
        ),
        params={
            "city":      {"type": "string",  "description": "City for weather check (default: Istanbul)"},
            "occasion":  {"type": "string",  "description": "Occasion: casual | work | formal | sport | date (default: casual)"},
            "style_pref":{"type": "string",  "description": "Style preferences (e.g. minimalist, colorful, classic)"},
        },
        package="you",
    )
    async def _outfit_of_day(city: str = "Istanbul", occasion: str = "casual", style_pref: str = ""):
        import urllib.request, urllib.parse, json
        # Get current weather
        try:
            geo_url = f"https://geocoding-api.open-meteo.com/v1/search?name={urllib.parse.quote(city)}&count=1"
            with urllib.request.urlopen(geo_url, timeout=8) as r:
                geo = json.loads(r.read())
            res = geo.get("results", [])
            if res:
                lat, lon = res[0]["latitude"], res[0]["longitude"]
                wurl = (f"https://api.open-meteo.com/v1/forecast?latitude={lat}&longitude={lon}"
                        "&current=temperature_2m,weathercode,precipitation_probability&timezone=auto")
                with urllib.request.urlopen(wurl, timeout=8) as r:
                    wdata = json.loads(r.read())
                cur = wdata["current"]
                temp  = cur["temperature_2m"]
                code  = cur["weathercode"]
                rain  = cur.get("precipitation_probability", 0)
                weather_summary = f"{temp}°C, {'rainy' if rain > 40 else 'dry'}, weather code {code}"
            else:
                weather_summary = "unknown weather"
        except Exception:
            weather_summary = "unknown weather"

        style_note = f" Style preference: {style_pref}." if style_pref else ""
        return {
            "city": city,
            "weather": weather_summary,
            "occasion": occasion,
            "instruction": (
                f"Suggest a complete outfit for someone in {city} today. "
                f"Current weather: {weather_summary}. Occasion: {occasion}.{style_note} "
                "Include: top, bottom (or dress), shoes, outerwear if needed, accessories. "
                "Be specific and practical. Give 2 outfit options."
            ),
        }

    # ── Recipe from Ingredients ───────────────────────────────────────────────
    @register(
        name="recipe_from_ingredients",
        description=(
            "Suggest recipes based on ingredients the user has. "
            "Use for: 'what can I cook with...', 'recipe with eggs and pasta', 'ne yemek yapayım', 'malzemelere göre yemek'."
        ),
        params={
            "ingredients": {"type": "string",  "description": "Comma-separated list of available ingredients", "required": True},
            "meal_type":   {"type": "string",  "description": "breakfast | lunch | dinner | snack | dessert (optional)"},
            "servings":    {"type": "integer", "description": "Number of servings (default 2)"},
            "dietary":     {"type": "string",  "description": "Dietary restrictions: vegetarian | vegan | gluten-free | etc."},
        },
        package="you",
    )
    async def _recipe_from_ingredients(ingredients: str, meal_type: str = "",
                                       servings: int = 2, dietary: str = ""):
        parts = [f"Ingredients available: {ingredients}"]
        if meal_type:
            parts.append(f"Meal type: {meal_type}")
        if dietary:
            parts.append(f"Dietary restrictions: {dietary}")
        parts.append(f"Servings: {servings}")
        return {
            "ingredients": ingredients,
            "meal_type":   meal_type or "any",
            "servings":    servings,
            "instruction": (
                f"Suggest 2-3 recipes using mainly these ingredients: {ingredients}. "
                + (f"Meal type: {meal_type}. " if meal_type else "")
                + (f"Dietary: {dietary}. " if dietary else "")
                + f"For {servings} servings. "
                "For each recipe include: name, ingredients with quantities, step-by-step instructions, and cooking time. "
                "Prioritize recipes where the user has most of the required ingredients."
            ),
        }

    # ── Activity Suggestions ──────────────────────────────────────────────────
    @register(
        name="activity_suggestions",
        description=(
            "Suggest activities or hobbies for free time. "
            "Use for: 'I'm bored', 'what should I do today', 'boşum ne yapayım', 'activity ideas', 'hobby suggestions'."
        ),
        params={
            "mood":        {"type": "string",  "description": "Current mood: energetic | relaxed | social | creative | adventurous"},
            "duration":    {"type": "string",  "description": "Available time: 30min | 1hr | 2hr | half-day | full-day"},
            "location":    {"type": "string",  "description": "Location context: home | outdoor | city | nature"},
            "interests":   {"type": "string",  "description": "User interests or hobbies (optional)"},
            "weather":     {"type": "string",  "description": "Current weather (optional, e.g. rainy, sunny)"},
        },
        package="you",
    )
    async def _activity_suggestions(mood: str = "relaxed", duration: str = "1hr",
                                    location: str = "home", interests: str = "",
                                    weather: str = ""):
        return {
            "mood": mood,
            "duration": duration,
            "location": location,
            "instruction": (
                f"Suggest 5 activities for someone who is {mood}, has {duration} available, "
                f"and is {'at home' if location == 'home' else 'in/near ' + location}. "
                + (f"Their interests: {interests}. " if interests else "")
                + (f"Current weather: {weather}. " if weather else "")
                + "For each activity: name, why it's good for this mood, what's needed, and how to get started right now. "
                "Mix familiar and new ideas. Make them feel genuinely excited to try."
            ),
        }

    # ── Personal Profile / FAQ ────────────────────────────────────────────────
    @register(
        name="get_my_profile",
        description=(
            "Return a summary of everything OZY2 knows about the user from memory. "
            "Use for: 'who am I', 'what do you know about me', 'my profile', 'SSS', 'tell me about myself'."
        ),
        params={
            "category": {"type": "string", "description": "Filter by: all | personal | preferences | goals | health | work (default: all)"},
        },
        package="you",
    )
    async def _get_my_profile(category: str = "all"):
        from pathlib import Path
        import json
        memory_file = Path.home() / ".ozy2" / "memory.json"
        if not memory_file.exists():
            return {
                "ok": False,
                "message": "No profile data yet. Tell OZY2 things about yourself and they'll be remembered!",
                "tip": "Try: 'Remember that I prefer dark mode', 'I work as a designer', 'My goal is to read 12 books this year'"
            }
        try:
            memories = json.loads(memory_file.read_text())
            if isinstance(memories, list):
                items = memories
            elif isinstance(memories, dict):
                items = list(memories.values()) if memories else []
            else:
                items = []

            if category != "all" and items:
                # Filter by category keyword
                items = [m for m in items if category.lower() in str(m).lower()]

            return {
                "ok": True,
                "category": category,
                "memory_count": len(items),
                "memories": items,
                "instruction": (
                    f"Based on these {len(items)} things OZY2 knows about the user, "
                    "create a personal profile/FAQ summary. "
                    "Organize by: Personal info, Preferences & tastes, Work & goals, Habits & routines. "
                    "Be warm and personal, as if introducing someone you know well. "
                    f"Data: {json.dumps(items, ensure_ascii=False)}"
                ),
            }
        except Exception as e:
            return {"ok": False, "error": str(e)}

    # ── Book Tracker ──────────────────────────────────────────────────────────
    @register(
        name="add_book",
        description=(
            "Add a book to the reading tracker. "
            "Use for: 'I started reading', 'add book', 'track this book', 'kitap ekle'."
        ),
        params={
            "title":    {"type": "string", "description": "Book title", "required": True},
            "author":   {"type": "string", "description": "Author name"},
            "status":   {"type": "string", "description": "reading | want_to_read | completed (default: reading)"},
            "total_pages": {"type": "integer", "description": "Total pages in the book"},
        },
        package="you",
    )
    async def _add_book(title: str, author: str = "", status: str = "reading",
                        total_pages: int = 0):
        from pathlib import Path
        from datetime import datetime
        import json
        db_file = Path.home() / ".ozy2" / "books.json"
        books = json.loads(db_file.read_text()) if db_file.exists() else []
        # Check duplicate
        existing = next((b for b in books if b["title"].lower() == title.lower()), None)
        if existing:
            return {"ok": False, "error": f"'{title}' is already in your library", "book": existing}
        book = {
            "id":           len(books) + 1,
            "title":        title,
            "author":       author,
            "status":       status,
            "total_pages":  total_pages,
            "current_page": 0,
            "added":        datetime.now().strftime("%Y-%m-%d"),
            "finished":     "",
            "rating":       0,
            "notes":        [],
        }
        books.append(book)
        db_file.parent.mkdir(parents=True, exist_ok=True)
        db_file.write_text(json.dumps(books, indent=2, ensure_ascii=False))
        return {"ok": True, "book": book, "total_books": len(books)}

    @register(
        name="update_reading_progress",
        description=(
            "Update reading progress for a book. "
            "Use for: 'I read 50 pages', 'I finished the book', 'I'm on page 200'."
        ),
        params={
            "title":        {"type": "string",  "description": "Book title (partial match ok)", "required": True},
            "current_page": {"type": "integer", "description": "Current page number"},
            "status":       {"type": "string",  "description": "reading | completed | paused"},
            "rating":       {"type": "integer", "description": "Rating 1-5 (set when completed)"},
        },
        package="you",
    )
    async def _update_reading_progress(title: str, current_page: int = 0,
                                       status: str = "", rating: int = 0):
        from pathlib import Path
        from datetime import datetime
        import json
        db_file = Path.home() / ".ozy2" / "books.json"
        if not db_file.exists():
            return {"ok": False, "error": "No books found. Add a book first."}
        books = json.loads(db_file.read_text())
        book = next((b for b in books if title.lower() in b["title"].lower()), None)
        if not book:
            return {"ok": False, "error": f"Book '{title}' not found"}
        if current_page:
            book["current_page"] = current_page
        if status:
            book["status"] = status
            if status == "completed" and not book["finished"]:
                book["finished"] = datetime.now().strftime("%Y-%m-%d")
        if rating:
            book["rating"] = max(1, min(5, rating))
        db_file.write_text(json.dumps(books, indent=2, ensure_ascii=False))
        progress = ""
        if book["total_pages"] and book["current_page"]:
            pct = int(book["current_page"] / book["total_pages"] * 100)
            progress = f"{pct}%"
        return {"ok": True, "book": book, "progress": progress}

    @register(
        name="add_book_note",
        description=(
            "Add a note or highlight to a book. "
            "Use for: 'note from this book', 'save this quote', 'kitaptan not al'."
        ),
        params={
            "title":   {"type": "string", "description": "Book title (partial match ok)", "required": True},
            "note":    {"type": "string", "description": "Note or quote to save", "required": True},
            "page":    {"type": "integer","description": "Page number (optional)"},
            "type":    {"type": "string", "description": "note | quote | highlight (default: note)"},
        },
        package="you",
    )
    async def _add_book_note(title: str, note: str, page: int = 0, type: str = "note"):
        from pathlib import Path
        from datetime import datetime
        import json
        db_file = Path.home() / ".ozy2" / "books.json"
        if not db_file.exists():
            return {"ok": False, "error": "No books found."}
        books = json.loads(db_file.read_text())
        book = next((b for b in books if title.lower() in b["title"].lower()), None)
        if not book:
            return {"ok": False, "error": f"Book '{title}' not found"}
        entry = {
            "type": type,
            "text": note,
            "page": page,
            "date": datetime.now().strftime("%Y-%m-%d"),
        }
        book.setdefault("notes", []).append(entry)
        db_file.write_text(json.dumps(books, indent=2, ensure_ascii=False))
        # Also save to memory for cross-context recall
        return {"ok": True, "book": book["title"], "note_saved": entry}

    @register(
        name="get_current_reading",
        description=(
            "Show books currently being read and reading stats. "
            "Use for: 'what am I reading', 'my books', 'reading progress', 'kitaplarım'."
        ),
        params={
            "status": {"type": "string", "description": "Filter: reading | completed | want_to_read | all (default: reading)"},
        },
        package="you",
    )
    async def _get_current_reading(status: str = "reading"):
        from pathlib import Path
        import json
        db_file = Path.home() / ".ozy2" / "books.json"
        if not db_file.exists():
            return {"ok": True, "books": [], "message": "No books tracked yet. Start by adding a book!"}
        books = json.loads(db_file.read_text())
        filtered = books if status == "all" else [b for b in books if b.get("status") == status]
        # Add progress percentage
        for b in filtered:
            if b.get("total_pages") and b.get("current_page"):
                b["progress_pct"] = int(b["current_page"] / b["total_pages"] * 100)
            else:
                b["progress_pct"] = 0
        stats = {
            "total": len(books),
            "reading": sum(1 for b in books if b.get("status") == "reading"),
            "completed": sum(1 for b in books if b.get("status") == "completed"),
            "want_to_read": sum(1 for b in books if b.get("status") == "want_to_read"),
        }
        return {"ok": True, "books": filtered, "stats": stats}

    # ── Smart Home (Home Assistant) ───────────────────────────────────────────
    @register(
        name="smarthome_status",
        description=(
            "Get the status of smart home devices via Home Assistant. "
            "Use for: 'home status', 'what lights are on', 'temperature at home', 'akıllı ev durumu'."
        ),
        params={
            "domain": {"type": "string", "description": "Filter by domain: light | switch | climate | sensor | all (default: all)"},
            "area":   {"type": "string", "description": "Filter by area/room name (optional)"},
        },
        package="you",
    )
    async def _smarthome_status(domain: str = "all", area: str = ""):
        from pathlib import Path
        import json, urllib.request
        cfg_file = Path.home() / ".ozy2" / "smarthome.json"
        if not cfg_file.exists():
            return {
                "ok": False,
                "setup_required": True,
                "message": "Smart home not configured. Go to Settings → Smart Home to connect Home Assistant.",
            }
        cfg = json.loads(cfg_file.read_text())
        ha_url   = cfg.get("url", "").rstrip("/")
        ha_token = cfg.get("token", "")
        if not ha_url or not ha_token:
            return {"ok": False, "setup_required": True, "message": "Home Assistant URL or token missing."}
        try:
            req = urllib.request.Request(
                f"{ha_url}/api/states",
                headers={"Authorization": f"Bearer {ha_token}", "Content-Type": "application/json"}
            )
            with urllib.request.urlopen(req, timeout=10) as r:
                states = json.loads(r.read())
            # Filter
            if domain != "all":
                states = [s for s in states if s["entity_id"].startswith(f"{domain}.")]
            if area:
                states = [s for s in states if area.lower() in s.get("attributes", {}).get("friendly_name", "").lower()]
            # Summarize
            summary = []
            for s in states[:50]:
                summary.append({
                    "entity":  s["entity_id"],
                    "name":    s.get("attributes", {}).get("friendly_name", s["entity_id"]),
                    "state":   s["state"],
                    "unit":    s.get("attributes", {}).get("unit_of_measurement", ""),
                })
            return {"ok": True, "devices": summary, "count": len(summary)}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    @register(
        name="smarthome_control",
        description=(
            "Control a smart home device via Home Assistant. "
            "Use for: 'turn on the lights', 'set thermostat to 22', 'lambayı aç', 'close the blinds'."
        ),
        params={
            "entity_id": {"type": "string", "description": "Entity ID (e.g. light.living_room) or friendly name", "required": True},
            "action":    {"type": "string", "description": "turn_on | turn_off | toggle | set_temperature | set_brightness", "required": True},
            "value":     {"type": "string", "description": "Value for set actions (e.g. temperature number or brightness 0-255)"},
        },
        package="you",
    )
    async def _smarthome_control(entity_id: str, action: str, value: str = ""):
        from pathlib import Path
        import json, urllib.request
        cfg_file = Path.home() / ".ozy2" / "smarthome.json"
        if not cfg_file.exists():
            return {"ok": False, "setup_required": True, "message": "Smart home not configured."}
        cfg = json.loads(cfg_file.read_text())
        ha_url   = cfg.get("url", "").rstrip("/")
        ha_token = cfg.get("token", "")
        if not ha_url or not ha_token:
            return {"ok": False, "error": "Home Assistant not configured."}

        domain = entity_id.split(".")[0] if "." in entity_id else "homeassistant"
        service_map = {
            "turn_on":  (domain, "turn_on"),
            "turn_off": (domain, "turn_off"),
            "toggle":   (domain, "toggle"),
            "set_temperature": ("climate", "set_temperature"),
            "set_brightness":  ("light", "turn_on"),
        }
        svc_domain, service = service_map.get(action, (domain, action))
        payload = {"entity_id": entity_id}
        if action == "set_temperature" and value:
            payload["temperature"] = float(value)
        elif action == "set_brightness" and value:
            payload["brightness"] = int(value)

        try:
            data = json.dumps(payload).encode()
            req = urllib.request.Request(
                f"{ha_url}/api/services/{svc_domain}/{service}",
                data=data,
                headers={
                    "Authorization": f"Bearer {ha_token}",
                    "Content-Type": "application/json",
                },
                method="POST",
            )
            with urllib.request.urlopen(req, timeout=10) as r:
                result = json.loads(r.read()) if r.read() else []
            return {"ok": True, "entity": entity_id, "action": action, "value": value}
        except Exception as e:
            return {"ok": False, "error": str(e)}
